import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import date

doc = Document()

# Page style
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

today = date.today().strftime('%Y年%m月%d日')

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return h

def add_para(text, bold=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align is not None:
        p.alignment = align
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

# ================================================================
# TITLE PAGE
# ================================================================

doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('A股科技板块投资分析报告')
run.font.name = 'Microsoft YaHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('生成日期：' + today)
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()
add_para('投资本金：6500元', size=11)
add_para('交易市场：A股主板（60xxxx / 00xxxx）', size=11)
add_para('选股方向：科技板块 - 半导体 / 电子元件 / 半导体设备', size=11)
add_para('分析工具：TradingAgents 多Agent系统 + mootdx 通达信数据', size=11)

doc.add_page_break()

# ================================================================
# SECTION 1
# ================================================================

add_heading_styled('一、筛选过程', level=1)

add_para('1.1 筛选条件', bold=True, size=12)
add_para('   板块限制：仅限主板（60xxxx / 00xxxx），排除创业板(300)、科创板(688)')
add_para('   价格区间：5元 ~ 65元（1手100股适配6500元本金）')
add_para('   行业范围：半导体、电子元器件、芯片设计/封测、光学光电子、IT/通信设备')
add_para('   趋势要求：均线多头排列、MACD金叉、RSI不过热（<80）、量能配合')

add_para('')
add_para('1.2 初筛候选池', bold=True, size=12)
add_para('从22只科技股中排除价格超范围的标的（如兆易创新445元、深南电路333元等），剩余21只进入趋势评分环节。评分维度包括：均线排列质量(0-7分)、多周期涨跌幅(5d/10d/20d/60d)、MACD方向与强度、RSI健康度、量比强弱等，满分22分。')

add_para('')
add_para('1.3 趋势动能评分排名（前10）', bold=True, size=12)

headers1 = ['排名', '代码', '名称', '价格', '趋势', '5日', '10日', '20日', '60日', 'RSI', '量比', '得分']
rows1 = [
    ['1', '600460', '士兰微', '32.59', '强多头', '+3.4%', '+9.9%', '+21.0%', '+6.6%', '73.4', '1.33', '21'],
    ['2', '603859', '能科科技', '46.34', '强多头', '+12.4%', '+12.6%', '+16.8%', '+0.7%', '71.4', '2.45', '21'],
    ['3', '002138', '顺络电子', '39.83', '强多头', '+4.3%', '+15.1%', '+12.4%', '+3.5%', '66.6', '1.52', '21'],
    ['4', '600703', '三安光电', '16.75', '强多头', '-0.5%', '+14.1%', '+25.6%', '+5.8%', '64.6', '0.81', '21'],
    ['5', '603005', '晶方科技', '36.80', '强多头', '+9.3%', '+16.6%', '+17.8%', '+15.4%', '72.1', '1.43', '21'],
    ['6', '603690', '至纯科技', '31.05', '强多头', '+12.8%', '+22.6%', '+21.2%', '+12.6%', '74.4', '2.40', '21'],
    ['7', '002273', '水晶光电', '39.70', '强多头', '+0.8%', '+19.6%', '+36.7%', '+46.5%', '74.3', '0.89', '20'],
    ['8', '002185', '华天科技', '15.88', '强多头', '+8.5%', '+16.6%', '+24.6%', '+15.1%', '78.6', '2.44', '19'],
    ['9', '600667', '太极实业', '13.69', '强多头', '+3.2%', '+34.7%', '+37.6%', '+45.0%', '78.4', '1.70', '19'],
    ['10', '000636', '风华高科', '34.00', '强多头', '+15.0%', '+37.0%', '+54.6%', '+61.4%', '81.9', '1.96', '19'],
]
add_table(headers1, rows1)

doc.add_page_break()

# ================================================================
# SECTION 2
# ================================================================

add_heading_styled('二、精选池深度对比（7只）', level=1)

add_para('结合量化筛选与主观判断，从初筛结果中挑选7只进入精选池进行深度对比分析。', size=11)

add_para('')
add_para('2.1 核心指标对比', bold=True, size=12)

headers2 = ['代码', '名称', '价格', '均线排列', '5日', '10日', '20日', 'RSI', 'MACD', '量比', '判定']
rows2 = [
    ['603690', '至纯科技', '31.05', '完美多头 7/7', '+12.8%', '+22.6%', '+21.2%', '74.4', '金叉', '2.40', '入选'],
    ['600460', '士兰微', '32.59', '完美多头 7/7', '+3.4%', '+9.9%', '+21.0%', '73.4', '金叉', '1.33', '入选'],
    ['002138', '顺络电子', '39.83', '多头 6/7', '+4.3%', '+15.1%', '+12.4%', '66.6', '金叉', '1.52', '备选'],
    ['000021', '深科技', '37.85', '完美多头 7/7', '+9.4%', '+23.2%', '+30.5%', '80.1', '金叉', '2.49', '超买'],
    ['603650', '彤程新材', '62.88', '完美多头 7/7', '+7.4%', '+12.0%', '+16.4%', '67.8', '金叉', '1.56', '价高'],
    ['000636', '风华高科', '34.00', '完美多头 7/7', '+15.0%', '+37.0%', '+54.6%', '81.9', '金叉', '1.96', '过热'],
    ['002484', '江海股份', '58.60', '完美多头 7/7', '+23.9%', '+44.9%', '+77.4%', '85.9', '金叉', '1.13', '过热'],
    ['603127', '昭衍新药', '38.88', '多头 5/7', '-1.5%', '-2.1%', '+5.3%', '55.7', '死叉', '0.86', '淘汰'],
]
add_table(headers2, rows2)

add_para('')
add_para('2.2 淘汰逻辑', bold=True, size=12)
add_para('  昭衍新药(603127)：MACD死叉 + 短期下跌 + 非科技板块（CRO医药），直接淘汰')
add_para('  江海股份(002484)：RSI 85.9极度超买，60日涨幅99.7%已翻倍，追高风险极大')
add_para('  风华高科(000636)：RSI 81.9严重超买，20日涨54.6%已大幅拉升，追入性价比差')
add_para('  彤程新材(603650)：指标优秀但62.88元/手几乎用满6500预算，无法分散')
add_para('  深科技(000021)：RSI 80.1超买，距20日高点-5.4%已有回撤，动能衰减')
add_para('  顺络电子(002138)：RSI 66.6最健康，均线多头6/7，但短期动能偏温和')

doc.add_page_break()

# ================================================================
# SECTION 3
# ================================================================

add_heading_styled('三、最终推荐：士兰微 + 至纯科技', level=1)

add_para('经逐只排查，最终选定两只：士兰微(600460) + 至纯科技(603690)。两只均为完美多头排列(7/7均线对齐)、MACD金叉、量能配合、细分互补。总成本约6364元，适配6500元本金。')

# --- 士兰微 ---
add_para('')
add_heading_styled('3.1 士兰微（600460）-- 功率半导体IDM龙头', level=2)

add_para('【基本面】', bold=True, size=11)
add_para('士兰微是国内少数实现IDM模式的功率半导体企业，覆盖IGBT、MOSFET、IPM等核心产品。半导体周期复苏中，功率器件是最先受益的环节之一。政策面持续催化（大基金三期 + 国产替代）。')

add_para('')
add_para('【技术面关键价位】', bold=True, size=11)
headers_slw = ['指标', '数值', '说明']
rows_slw = [
    ['当前价', '32.59元', '2026-05-20收盘价'],
    ['MA5', '31.66元', '5日均线，短期支撑'],
    ['MA10', '30.97元', '10日均线，中期支撑'],
    ['MA20', '29.53元', '20日均线，趋势生命线'],
    ['MA60', '28.53元', '60日均线，长期趋势'],
    ['布林上轨', '33.00元', '距现价仅+1.3%，突破即加速'],
    ['20日最高', '32.98元', '短期阻力位'],
    ['ATR(14)', '1.27 (3.9%)', '日均波动约1.27元'],
    ['RSI(14)', '73.4', '偏热但未极端'],
]
add_table(headers_slw, rows_slw)

add_para('')
add_para('【近期走势】', bold=True, size=11)
add_para('04/27-05/08在28-30区间横盘整理，05/11放量突破30后回踩确认(05/15收30.95)，05/18-05/20连阳拉升(31.05-32.35-32.59)。走势扎实稳健，回踩有支撑、突破有量，不是急拉虚高。')

add_para('')
add_para('【交易计划】', bold=True, size=11)
headers_plan1 = ['项目', '方案A（推荐）', '方案B（保守）', '方案C（激进）']
rows_plan1 = [
    ['买入区间', '31.00-31.70元', '30.50-31.00元', '31.70-32.00元'],
    ['买入逻辑', '回踩MA5-MA10', '深度回调至MA10以下', '回踩MA5浅回调'],
    ['距现价', '-2.7%至-4.9%', '-4.9%至-6.4%', '-1.8%至-2.7%'],
    ['止损价', '29.45元', '29.45元', '29.45元'],
    ['止损幅度', '-5.0%', '-5.0%', '-5.0%'],
    ['止损逻辑', 'MA20(29.53)+5%铁律', 'MA20+5%铁律', 'MA20+5%铁律'],
    ['止盈T1', '34.10元 (+10%)', '34.10元 (+10%)', '34.10元 (+10%)'],
    ['止盈T2', '35.20元 (+13.5%)', '35.20元 (+13.5%)', '35.20元 (+13.5%)'],
    ['盈亏比', '2.0 : 1', '2.0 : 1', '2.0 : 1'],
    ['1手成本', '3100-3170元', '3050-3100元', '3170-3200元'],
]
add_table(headers_plan1, rows_plan1)

# --- 至纯科技 ---
add_para('')
add_heading_styled('3.2 至纯科技（603690）-- 半导体清洗设备', level=2)

add_para('【基本面】', bold=True, size=11)
add_para('至纯科技是国内半导体清洗设备龙头，产品覆盖槽式清洗、单片清洗等核心工艺环节。随着国内晶圆厂扩产（中芯国际、华虹等），清洗设备需求持续增长。国产替代率仍低，成长空间大。')

add_para('')
add_para('【技术面关键价位】', bold=True, size=11)
headers_zc = ['指标', '数值', '说明']
rows_zc = [
    ['当前价', '31.05元', '2026-05-20收盘价'],
    ['MA5', '29.13元', '5日均线，短线支撑'],
    ['MA10', '27.79元', '10日均线，中线支撑'],
    ['MA20', '26.74元', '20日均线，趋势生命线'],
    ['MA60', '26.31元', '60日均线，长线支撑'],
    ['布林上轨', '30.00元', '现价已突破上轨！'],
    ['20日最高', '31.82元', '短期阻力位'],
    ['ATR(14)', '1.50 (4.8%)', '日均波动约1.50元'],
    ['RSI(14)', '74.4', '偏热，短线需回调消化'],
]
add_table(headers_zc, rows_zc)

add_para('')
add_para('【近期走势】', bold=True, size=11)
add_para('04/27-05/08底部盘整24-26区间，05/11-05/14放量突破拉升(26.50-28.63)，05/15-05/18冲高回落洗盘(28.57-28.38)，05/19-05/20再度爆量上攻(29.04-31.05，单日+6.9%)。呈现"拉升-洗盘-再拉升"的强势节奏。但当前已突破布林上轨，短线严重过热，必须等待回调。')

add_para('')
add_para('【交易计划】', bold=True, size=11)
headers_plan2 = ['项目', '方案A（推荐）', '方案B（激进）', '方案C（保守）']
rows_plan2 = [
    ['买入区间', '28.80-29.80元', '30.00-31.00元', '27.80-28.50元'],
    ['买入逻辑', '回踩MA5附近', '追突破（怕踏空）', '深调至MA10'],
    ['距现价', '-4.0%至-7.2%', '0%至-3.4%', '-8.2%至-10.5%'],
    ['止损价', '27.36元', '28.50元', '27.36元'],
    ['止损幅度', '-5.0%', '-5.0%', '-5.0%'],
    ['止损逻辑', 'MA10+5%铁律', '5%铁律', 'MA10+5%铁律'],
    ['止盈T1', '31.68元 (+10%)', '33.00元 (+10%)', '30.58元 (+10%)'],
    ['止盈T2', '33.12元 (+15%)', '34.50元 (+15%)', '31.97元 (+15%)'],
    ['盈亏比', '2.0 : 1', '2.0 : 1', '2.0 : 1'],
    ['1手成本', '2880-2980元', '3000-3100元', '2780-2850元'],
]
add_table(headers_plan2, rows_plan2)

doc.add_page_break()

# ================================================================
# SECTION 4
# ================================================================

add_heading_styled('四、组合方案与仓位分配', level=1)

add_para('')
add_para('4.1 推荐组合（总成本 5960-6150元）', bold=True, size=12)

headers_combo = ['标的', '代码', '数量', '预计成本', '占总资金', '细分赛道']
rows_combo = [
    ['士兰微', '600460', '1手（100股）', '3100-3170元', '47.7%-48.8%', '功率半导体IDM'],
    ['至纯科技', '603690', '1手（100股）', '2880-2980元', '44.3%-45.8%', '半导体清洗设备'],
    ['剩余现金', '', '', '350-520元', '5.4%-8.0%', '备用/补仓'],
]
add_table(headers_combo, rows_combo)

add_para('')
add_para('4.2 操作节奏', bold=True, size=12)
add_para('  第一步：士兰微回踩31.00-31.70区间，买入1手（约3150元）')
add_para('  第二步：至纯科技回踩28.80-29.80区间，买入1手（约2930元）')
add_para('  注意：两只不一定同时到达买点，哪个先到先买哪个，不要追高')
add_para('  若至纯不回踩直接拉升，放弃至纯，士兰微单只满仓也是合理选择')

add_para('')
add_para('4.3 退出纪律', bold=True, size=12)
add_para('  止损铁律：任何一只跌幅达-5%，不问原因，立即卖出')
add_para('  止盈纪律：到+8%卖一半，锁定利润；到+10%以上分批清仓')
add_para('  时间止损：持仓超过7个交易日未达目标，重新评估是否离场')
add_para('  禁止操作：不补仓、不做T、不满仓一只票、每周最多交易2次')

doc.add_page_break()

# ================================================================
# SECTION 5
# ================================================================

add_heading_styled('五、风险提示', level=1)

add_para('')
add_para('5.1 市场风险', bold=True, size=12)
add_para('  半导体板块短期累计涨幅较大，存在板块性回调风险')
add_para('  两只标的RSI均在73-74区间，处于"偏热"状态，对利空敏感')
add_para('  至纯科技突破布林上轨，短期超买程度高于士兰微')

add_para('')
add_para('5.2 个股风险', bold=True, size=12)
add_para('  士兰微：功率半导体竞争加剧，IGBT价格承压；公司盈利能力波动较大')
add_para('  至纯科技：流通市值偏小，日均成交额偏低，存在流动性风险；客户集中度高')

add_para('')
add_para('5.3 操作风险', bold=True, size=12)
add_para('  追高买入是最大风险 -- 严格按买入区间执行，不追涨')
add_para('  不止损是第二大风险 --  -5%铁律不可妥协')
add_para('  同时持仓两只半导体股，板块系统性风险无法分散')

add_para('')
add_para('')
add_para('-' * 50)
add_para('')
p = doc.add_paragraph()
run = p.add_run('免责声明')
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
add_para('本报告由AI辅助生成，仅供学习研究与个人决策参考，不构成任何投资建议。投资有风险，入市需谨慎。投资者应独立判断并承担投资风险。', size=9, color=(0x99, 0x99, 0x99))

# Save
output_path = 'C:/Users/zhzsh/Desktop/A股科技板块投资分析报告.docx'
doc.save(output_path)
print('Done: ' + output_path)
