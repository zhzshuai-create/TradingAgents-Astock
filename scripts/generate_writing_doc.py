"""生成六级作文模板 + 素材 Word 文档"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── 标题 ──
title = doc.add_heading("CET-6 六级作文万能模板 + 素材库", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("考场用时：30分钟  |  目标分数：12-15分  |  字数要求：150-200词").alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("─" * 60)

# ═══════════════════════ part 1 模板 ═══════════════════════
doc.add_heading("一、万能模板（填空即用）", level=1)

template_text = """There is no denying that the value of (TOPIC) has become virtually impossible to ignore. From my perspective, (TOPIC) serves as a driving force behind long-term personal growth and development.

The importance of (TOPIC) can be understood from two angles. Above all, it exerts a profound impact on (AREA). Specifically, (TOPIC) equips individuals with the essential ability to (ABILITY). Equally significant is its contribution to (AREA2). More precisely, those who are armed with (TOPIC) are far more likely to (RESULT). Taken together, these factors are living proof of the significance of (TOPIC).

All in all, (TOPIC) is an indispensable component of our life, which holds the potential to bring far-reaching rewards. Only by embracing (TOPIC) can we unlock its full potential and lead a more rewarding life."""

p = doc.add_paragraph()
run = p.add_run(template_text)
run.font.name = "Consolas"
run.font.size = Pt(10)

# ── 填空说明 ──
doc.add_heading("填空说明", level=3)
table = doc.add_table(9, 2, style="Light Grid Accent 1")
instructions = [
    ("占位符", "填什么"),
    ("(TOPIC)", "主题词（如：teamwork / perseverance / critical thinking）"),
    ("(AREA)", "第一个受益领域（如：academic life / school life）"),
    ("(ABILITY)", "培养什么能力（如：communicate effectively / solve problems）"),
    ("(AREA2)", "第二个受益领域（如：career development / future workplace）"),
    ("(RESULT)", "带来什么结果（如：achieve success / overcome obstacles）"),
    ("❌ 注意", "所有 (TOPIC) 必须用同一个词，不能变来变去"),
    ("❌ 注意", "不要照抄原词，根据题目要求替换"),
    ("❌ 注意", "字数不够时：第二段在 (ABILITY) 和 (RESULT) 之间插入例子"),
]
for i, (a, b) in enumerate(instructions):
    table.cell(i, 0).text = a
    table.cell(i, 1).text = b

doc.add_paragraph()

# ═══════════════════════ part 2 实例 ═══════════════════════
doc.add_heading("二、完整示例（题目：The Importance of Teamwork）", level=1)

example_text = """There is no denying that the value of teamwork has become virtually impossible to ignore. From my perspective, teamwork serves as a driving force behind long-term personal growth and development.

The importance of teamwork can be understood from two angles. Above all, it exerts a profound impact on academic life. Specifically, teamwork equips individuals with the essential ability to communicate and collaborate effectively. Equally significant is its contribution to career development. More precisely, those who are armed with teamwork skills are far more likely to achieve common goals and earn the trust of colleagues. Taken together, these factors are living proof of the significance of teamwork.

All in all, teamwork is an indispensable component of our life, which holds the potential to bring far-reaching rewards. Only by embracing teamwork can we unlock its full potential and lead a more rewarding life."""

p = doc.add_paragraph()
run = p.add_run(example_text)
run.font.name = "Consolas"
run.font.size = Pt(10)

doc.add_paragraph()

# ═══════════════════════ part 3 素材 ═══════════════════════
doc.add_heading("三、5 个万能素材（直接背，考场拿来用）", level=1)

materials = [
    ("素材① — 大学生实习经历",
     "Take the experience of many college students today as an example. Those who actively seek internships during their studies not only gain hands-on skills but also develop a clearer understanding of their career direction. In contrast, students who focus solely on textbooks often find themselves unprepared when entering the job market. This contrast vividly illustrates why practical experience matters as much as academic knowledge.",
     "preparation, practice, employment, ability, planning"),

    ("素材② — 马云创业多次失败",
     'A well-known example is Jack Ma, the founder of Alibaba. Before achieving success, he was rejected by dozens of companies, including KFC, and failed in his first two business ventures. Yet it was precisely these setbacks that shaped his resilience and ultimately led to one of the world\'s largest e-commerce platforms. His story serves as a powerful reminder that failure is not the end, but a necessary step toward success.',
     "perseverance, challenge, innovation, failure, attitude"),

    ("素材③ — 疫情在线教育普及",
     "The COVID-19 pandemic brought an unexpected transformation to education. When schools were forced to close, millions of students turned to online platforms such as Zoom and Tencent Meeting. This shift, though initially difficult, demonstrated that learning can happen anywhere — not just in traditional classrooms — and that those who adapt quickly to new technology gain a lasting advantage.",
     "technology, adaptation, innovation, learning, change"),

    ("素材④ — 碳中和政策",
     "China's carbon neutrality policy, which aims to peak carbon emissions by 2030 and achieve net-zero by 2060, is a telling example of long-term planning. Rather than pursuing short-term economic gains at the expense of the environment, the government has chosen a path that balances development with sustainability. This policy reflects a profound sense of responsibility not only to the present generation but to those yet to come.",
     "environment, responsibility, long-term planning, balance"),

    ("素材⑤ — AI 取代重复工作",
     "The rapid advancement of artificial intelligence is reshaping the job market at an unprecedented pace. Routine and repetitive tasks, from data entry to assembly-line work, are increasingly being handled by machines. Those who rely solely on a single skill set risk being left behind, while those who commit to lifelong learning are better positioned to thrive in an AI-driven era. This trend highlights the urgent need for continuous self-improvement.",
     "technology, lifelong learning, competition, self-improvement"),
]

for title_text, body, tags in materials:
    doc.add_heading(title_text, level=2)
    p = doc.add_paragraph()
    run = p.add_run(body)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    p = doc.add_paragraph()
    p.add_run("可套话题：").bold = True
    p.add_run(f"  {tags}")

# ═══════════════════════ part 4 速查表 ═══════════════════════
doc.add_heading("四、速查表：看到题目关键词 → 选哪个素材", level=1)

lookup = [
    ("关键词", "素材"),
    ("preparation / planning / practice / ability", "① 实习经历"),
    ("failure / challenge / persistence / attitude", "② 马云创业"),
    ("technology / online / change / adaptation", "③ 在线教育"),
    ("environment / responsibility / future / balance", "④ 碳中和"),
    ("AI / competition / learning / self-improvement", "⑤ AI 替代"),
]
t = doc.add_table(7, 2, style="Light Grid Accent 1")
for i, (k, v) in enumerate(lookup):
    t.cell(i, 0).text = k
    t.cell(i, 1).text = v

doc.add_paragraph()

# ═══════════════════════ part 5 使用技巧 ═══════════════════════
doc.add_heading("五、考场快速套用技巧", level=1)

tips = [
    "① 拿到题目 → 圈出主题词 → 替换模板中的 (TOPIC)",
    "② 看题目属于哪个方向（能力？坚持？科技？）→ 查速查表 → 挑对应的素材",
    "③ 在模板「(ABILITY) 和 (RESULT)」句子后，用「For instance,」开头插入素材",
    "④ 所有 (TOPIC) 用同一个词，不要一会用 teamwork 一会用 cooperation",
    "⑤ 如果字数不够，在第二段中间（素材之后再补一句 More significantly...）",
    "⑥ (TOPIC) 都是名词，如果题目是动词短语（How to...），先转成名词",
    "   例：How to manage time → time management",
    "⑦ 结尾句不要写错：Only by...can we...（倒装，by 后面接 doing）",
]
for tip in tips:
    doc.add_paragraph(tip, style="List Bullet")

# ── Footer ──
doc.add_paragraph("─" * 60)
p = doc.add_paragraph()
p.add_run("Created for CET-6 Writing Preparation  |  2026-06").font.size = Pt(8)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Save ──
path = r"C:\Users\zhzsh\Desktop\CET6_作文模板_素材库.docx"
doc.save(path)
print(f"✅ 已保存: {path}")
